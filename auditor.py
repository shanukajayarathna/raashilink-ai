from docx import Document
import re
import os

def audit():
    try:
        doc = Document('report.docx')
    except Exception as e:
        print(f"Error: {e}")
        return

    full_text = []
    placeholders = []
    facts = {"matches_cold": 9247, "matches_avg": 3535, "horoscope": 2994, "my_chart": 587, "health": 3, "login": 441}

    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            full_text.append(t)
            if re.search(r"INSERT|TBD|XX|\[.*?\]", t, re.IGNORECASE): placeholders.append(t)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text.strip()
                if t:
                    full_text.append(t)
                    if re.search(r"INSERT|TBD|XX|\[.*?\]", t, re.IGNORECASE): placeholders.append(t)

    content = "\n".join(full_text)
    
    print("\n[Section 1: Placeholders]")
    if placeholders:
        print(f"Found {len(placeholders)} matches.")
        for p in list(set(placeholders))[:5]: print(f"- {p}")
    else: print("None found.")

    print("\n[Section 2: Performance Check]")
    for key, val in facts.items():
        count = content.count(str(val))
        print(f"{key} ({val}): {'FOUND (' + str(count) + ' times)' if count > 0 else 'MISSING'}")

    print("\n[Section 3: Weak Spots]")
    checks = {
        "GDPR/Privacy": ["gdpr", "privacy", "protection"],
        "Security": ["security", "encryption", "jwt", "https"],
        "Methodology": ["agile", "scrum", "sdlc", "waterfall"],
        "Testing": ["unit", "integration", "cypress", "jest"],
        "Limitations": ["limitation", "bottleneck"],
        "Ethics": ["ethic", "consent"],
        "Deployment": ["docker", "ci/cd", "vercel"]
    }
    for cat, keys in checks.items():
        found = any(k in content.lower() for k in keys)
        print(f"{cat}: {'PASS' if found else 'FAIL'}")

    informal = [p for p in full_text if any(f" {w} " in p.lower() for w in ["don't", "can't", "basically", "very"])]
    print(f"\n[Section 4: Writing] Informal snippets found: {len(informal)}")

audit()
